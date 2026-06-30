from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "deliverables/Problem_Statement_Solution_Explainer.docx"

BLUE = "1F4D78"
MID_BLUE = "2E74B5"
INK = "1F2937"
MUTED = "5B677A"
LIGHT_BLUE = "E8F1FA"
LIGHT_GRAY = "F4F6F8"
BORDER = "C9D3DF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_run(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=16 if level == 1 else 13, color=MID_BLUE if level == 1 else BLUE, bold=True)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2)
    else:
        r = p.add_run(text)
        set_run(r)


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, "B8CADB")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run(r, size=10.5, color=BLUE, bold=True)
    r = p.add_run(body)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_solution_table(doc):
    rows = [
        ("Fast checkout", "POS cart, payment method selection, ID verification prompt, receipt flow"),
        ("Inventory control", "Products, stock quantities, packaging units, categories, and printable tags"),
        ("Customer history", "Customer records built from sales, contact details, and order counts"),
        ("Operational visibility", "Dashboard, reports, revenue totals, payment mix, and recent orders"),
        ("Secure access", "Google login with Firebase plus email/password flows and reset support"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Capability"
    hdr[1].text = "How the solution responds"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                set_run(r, bold=True, color=BLUE)
    for capability, response in rows:
        cells = table.add_row().cells
        cells[0].text = capability
        cells[1].text = response
        for i, cell in enumerate(cells):
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(r, size=10.5, color=INK, bold=(i == 0))


def add_diagram(doc):
    add_heading(doc, "Small Diagram: Solution Flow", 2)
    table = doc.add_table(rows=2, cols=5)
    set_table_width(table)
    labels = [
        "Store Staff",
        "RetailFlow POS",
        "Inventory + Orders",
        "Customers + Receipts",
        "Reports + Decisions",
    ]
    details = [
        "Logs in and scans or selects products",
        "Builds bill, verifies sale, captures payment",
        "Updates product stock and transaction history",
        "Stores buyer context and prints receipts/tags",
        "Shows sales, stock, trends, and next actions",
    ]
    widths = [1700, 1900, 1950, 1950, 1860]
    for row_idx, values in enumerate((labels, details)):
        row = table.rows[row_idx]
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.width = Pt(widths[i] / 20)
            set_cell_shading(cell, LIGHT_BLUE if row_idx == 0 else "FFFFFF")
            set_cell_border(cell, "9FB8D0")
            set_cell_margins(cell, top=120, bottom=120, start=100, end=100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    set_run(r, size=9.5 if row_idx == 1 else 10, color=BLUE if row_idx == 0 else INK, bold=(row_idx == 0))
    add_para(doc, "Flow: Store Staff -> POS -> Inventory/Orders -> Customers/Receipts -> Reports/Decisions", size=9.5, color=MUTED, italic=True, after=8)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Problem Statement & Solution Explainer")
    set_run(r, size=22, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("RetailFlow / BarrelFlow POS for retail and liquor-store operations")
    set_run(r, size=11.5, color=MUTED)

    add_callout(
        doc,
        "Purpose:",
        "This brief explains the business problem and how the proposed POS product solves it for small retail stores that need faster billing, cleaner inventory control, and clearer operational reporting.",
    )

    add_heading(doc, "Problem Statement", 1)
    add_para(
        doc,
        "Small retail and liquor stores often run daily operations across disconnected tools: manual billing, spreadsheet inventory, informal customer tracking, and paper receipts. This creates slow checkout lines, inaccurate stock counts, limited sales visibility, and avoidable mistakes during busy store hours.",
        after=8,
    )
    add_para(doc, "The core problem is operational fragmentation:", bold=True, color=BLUE, after=4)
    add_bullet(doc, "Billing is slow when cashiers must manually calculate carts, discounts, taxes, and payment status.")
    add_bullet(doc, "Inventory is unreliable when product sales do not immediately update stock and packaging quantities.")
    add_bullet(doc, "Customer and order history are hard to retrieve when receipts, phone numbers, and previous purchases are spread across multiple places.")
    add_bullet(doc, "Store owners lack real-time visibility into revenue, best-selling items, payment mix, and low-stock risks.")
    add_bullet(doc, "Authentication and staff access need to be simple enough for store teams but controlled enough for business use.")

    add_heading(doc, "Solution Overview", 1)
    add_para(
        doc,
        "RetailFlow / BarrelFlow is a modern web-based POS workspace that brings billing, inventory, orders, customers, promotions, settings, and reporting into one interface. The product is built as a Next.js application with a polished cashier dashboard and Google authentication support.",
        after=8,
    )
    add_solution_table(doc)

    add_heading(doc, "How the Product Works", 1)
    add_bullet(doc, "Cashiers sign in, open the POS terminal, select or scan products, verify age when required, choose payment mode, and confirm the sale.")
    add_bullet(doc, "Inventory screens let the store maintain products, categories, stock quantities, pricing, packaging units, and printable product tags.")
    add_bullet(doc, "Orders, customers, and reports give owners a structured view of transactions, customer activity, sales totals, and operational trends.")
    add_bullet(doc, "Settings allow the store to manage branding, receipt content, theme, employee details, and business preferences.")

    add_diagram(doc)

    add_heading(doc, "Current State and Next Step", 1)
    add_para(
        doc,
        "The current product is a strong working POS prototype with authentication and a complete front-end operating flow. Its business data layer currently relies heavily on a browser-local mock API for products, orders, customers, and reports. The recommended next step is to connect those modules to real Supabase database tables with row-level security so each store account has durable, isolated production data.",
        after=8,
    )
    add_callout(
        doc,
        "Outcome:",
        "The solution turns fragmented store operations into one faster, clearer workflow: sell products, update inventory, track customers, print receipts, and review performance from the same POS system.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("RetailFlow / BarrelFlow POS - Problem & Solution Brief")
    set_run(r, size=9, color=MUTED)

    doc.save(OUT)


if __name__ == "__main__":
    build()
